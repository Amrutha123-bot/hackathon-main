#read the pdf
#extract text
#clean text

#input is the insurance pdf file
#the output is the raw text file

import os
import logging
from datetime import datetime
# from pydoc import doc

# from langchain_community.document_loaders import (TextLoader, UnstructuredPDFLoader, UnstructuredWordDocumentLoader)
# from langchain_community.document_loaders import (
#     TextLoader,
#     PyPDFLoader,
#     Docx2txtLoader
# )
# from pypdf import PdfReader
# from docx import Document as DocxDocument
# from langchain_core.documents import Document
# from dataclasses import dataclass
# from typing import List
# from langchain_core.documents import Document
# logger = logging.getLogger(__name__)

# @dataclass
# class DocumentLoadResult:
#     documents: List[Document]
#     failed_files: List[str]
#     total_loaded: int

# class DocumentLoaderService:
#     # def get_loader(self, file_path: str):
#     #  #unstructured pdf loader#the word doc loader#the text loader
#     #     extension = os.path.splitext(file_path)[1].lower()
#     #     if(extension == '.pdf'):
#     #         return PyPDFLoader(file_path)
#     #     elif extension == '.docx':
#     #         return Docx2txtLoader(file_path)
#     #     elif extension == '.txt':
#     #         return TextLoader(file_path)
#     #     logger.warning(f"Unsupported file type: {extension}. ")
#     #     return None

#     def load_documents(
#         self,
#         directory_path: str,
#         strict_mode: bool = False
#     ) -> DocumentLoadResult:

#         documents = []
#         failed_files = []

#         if not os.path.exists(directory_path):
#             raise FileNotFoundError(
#                 f"Directory not found: {directory_path}"
#             )

#         for root, dirs, files in os.walk(directory_path):

#             for file in files:

#                 file_path = os.path.join(root, file)
#                 extension = os.path.splitext(file)[1].lower()

#                 try:

#                     # ---------------- PDF ----------------

#                     if extension == ".pdf":

#                         reader = PdfReader(file_path)

#                         text = ""

#                         for page in reader.pages:
#                             text += page.extract_text() or ""

#                     # ---------------- DOCX ----------------

#                     elif extension == ".docx":

#                         doc = DocxDocument(file_path)

#                         text = "\n".join(
#                             para.text
#                             for para in doc.paragraphs
#                         )

#                     # ---------------- TXT ----------------

#                     elif extension == ".txt":

#                         with open(
#                             file_path,
#                             "r",
#                             encoding="utf-8"
#                         ) as f:

#                             text = f.read()

#                     # ---------------- Unsupported ----------------

#                     else:

#                         logger.warning(
#                             f"Unsupported file type: {extension}"
#                         )

#                         failed_files.append(file_path)
#                         continue

#                     # Create LangChain Document

#                     document = Document(
#                         page_content=text,
#                         metadata={
#                             "source": file,
#                             "file_type": extension,
#                             "file_path": file_path,
#                             "loaded_at": datetime.now().isoformat(),
#                         },
#                     )

#                     documents.append(document)

#                 except Exception as e:

#                     logger.error(
#                         f"Error loading file {file_path}: {e}"
#                     )

#                     failed_files.append(file_path)

#                     if strict_mode:
#                         raise

#         return DocumentLoadResult(
#             documents=documents,
#             failed_files=failed_files,
#             total_loaded=len(documents),
#         )
#to test the module 1 code
# pdf_service = PDFService()
# result = pdf_service.load_documents('path/to/your/directory', strict_mode=False)
# print(result.total_loaded)
# print(result.failed_files)

# for doc in result.documents:
#     print(doc.metadata)

import os
import logging
from datetime import datetime
from dataclasses import dataclass
from typing import List

from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


@dataclass
class DocumentLoadResult:
    documents: List[Document]
    failed_files: List[str]
    total_loaded: int


class DocumentLoaderService:

    def load_documents(
        self,
        directory_path: str,
        strict_mode: bool = False
    ) -> DocumentLoadResult:

        documents = []
        failed_files = []

        if not os.path.exists(directory_path):
            raise FileNotFoundError(
                f"Directory not found: {directory_path}"
            )

        for root, dirs, files in os.walk(directory_path):

            for file in files:

                file_path = os.path.join(root, file)
                extension = os.path.splitext(file)[1].lower()

                try:

                    # =====================================================
                    # PDF
                    # =====================================================
                    if extension == ".pdf":

                        reader = PdfReader(file_path)

                        for page_num, page in enumerate(reader.pages):

                            text = page.extract_text() or ""

                            # Skip completely empty pages
                            if not text.strip():
                                continue

                            document = Document(
                                page_content=text,
                                metadata={
                                    "source": file,
                                    "page": page_num + 1,
                                    "file_type": extension,
                                    "file_path": file_path,
                                    "loaded_at": datetime.now().isoformat(),
                                    
                                },
                            )

                            documents.append(document)

                    # =====================================================
                    # DOCX
                    # =====================================================
                    elif extension == ".docx":

                        doc = DocxDocument(file_path)

                        text = "\n".join(
                            para.text
                            for para in doc.paragraphs
                        )

                        document = Document(
                            page_content=text,
                            metadata={
                                "source": file,
                                "page": 1,
                                "file_type": extension,
                                "file_path": file_path,
                                "loaded_at": datetime.now().isoformat(),
                            },
                        )

                        documents.append(document)

                    # =====================================================
                    # TXT
                    # =====================================================
                    elif extension == ".txt":

                        with open(
                            file_path,
                            "r",
                            encoding="utf-8"
                        ) as f:

                            text = f.read()

                        document = Document(
                            page_content=text,
                            metadata={
                                "source": file,
                                "page": 1,
                                "file_type": extension,
                                "file_path": file_path,
                                "loaded_at": datetime.now().isoformat(),
                            },
                        )

                        documents.append(document)

                    # =====================================================
                    # Unsupported
                    # =====================================================
                    else:

                        logger.warning(
                            f"Unsupported file type: {extension}"
                        )

                        failed_files.append(file_path)

                except Exception as e:

                    logger.error(
                        f"Error loading file {file_path}: {e}"
                    )

                    failed_files.append(file_path)

                    if strict_mode:
                        raise

        logger.info(f"Successfully loaded {len(documents)} documents.")

        return DocumentLoadResult(
            documents=documents,
            failed_files=failed_files,
            total_loaded=len(documents),
        )

    def load_file(
    self,
    file_path: str,
    strict_mode: bool = False
) -> DocumentLoadResult:

        documents = []
        failed_files = []

        if not os.path.exists(file_path):
            raise FileNotFoundError(
                f"File not found: {file_path}"
            )

        file = os.path.basename(file_path)
        extension = os.path.splitext(file)[1].lower()

        try:

            # =====================================================
            # PDF
            # =====================================================
            if extension == ".pdf":

                reader = PdfReader(file_path)

                for page_num, page in enumerate(reader.pages):

                    text = page.extract_text() or ""

                    if not text.strip():
                        continue

                    document = Document(
                        page_content=text,
                        metadata={
                            "source": file,
                            "page": page_num + 1,
                            "file_type": extension,
                            "file_path": file_path,
                            "loaded_at": datetime.now().isoformat(),
                        },
                    )

                    documents.append(document)

            # =====================================================
            # DOCX
            # =====================================================
            elif extension == ".docx":

                doc = DocxDocument(file_path)

                text = "\n".join(
                    para.text
                    for para in doc.paragraphs
                )

                if text.strip():

                    document = Document(
                        page_content=text,
                        metadata={
                            "source": file,
                            "page": 1,
                            "file_type": extension,
                            "file_path": file_path,
                            "loaded_at": datetime.now().isoformat(),
                        },
                    )

                    documents.append(document)

            # =====================================================
            # TXT
            # =====================================================
            elif extension == ".txt":

                with open(
                    file_path,
                    "r",
                    encoding="utf-8"
                ) as f:

                    text = f.read()

                if text.strip():

                    document = Document(
                        page_content=text,
                        metadata={
                            "source": file,
                            "page": 1,
                            "file_type": extension,
                            "file_path": file_path,
                            "loaded_at": datetime.now().isoformat(),
                        },
                    )

                    documents.append(document)

            # =====================================================
            # Unsupported
            # =====================================================
            else:

                logger.warning(
                    f"Unsupported file type: {extension}"
                )

                failed_files.append(file_path)

        except Exception as e:

            logger.error(
                f"Error loading file {file_path}: {e}"
            )

            failed_files.append(file_path)

            if strict_mode:
                raise

        logger.info(
            f"Loaded {len(documents)} documents from {file}"
        )

        return DocumentLoadResult(
            documents=documents,
            failed_files=failed_files,
            total_loaded=len(documents),
        )